from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("吵架记录表.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE8")


def configure_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("吵架记录表")
set_run_font(run, size=20, bold=True, color="22324A")

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note_run = note.add_run("用于复盘事实、压力、身体状态和沟通质量，不用于指责。")
set_run_font(note_run, size=10, color="5A6578")

headers = [
    "吵架时间",
    "大致内容",
    "七天内的\n压力性事件",
    "3天内\n生理状态",
    "3天内\n睡眠情况",
    "3天内环境和\n人际关系",
    "沟通情况",
    "沟通满意度\n（1到10分）",
    "大致结果",
]

table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
set_table_borders(table)

widths = [1350, 2500, 2300, 1900, 1900, 2300, 2100, 1500, 1900]
configure_table_width(table, widths)

header_row = table.rows[0]
header_row.height = Cm(1.05)
header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
set_repeat_table_header(header_row)

for i, text in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_shading(cell, "EEF3F8")
    set_cell_margins(cell, top=120, start=90, bottom=120, end=90)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, bold=True, color="22324A")

for _ in range(7):
    row = table.add_row()
    row.height = Cm(2.05)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for idx, cell in enumerate(row.cells):
        set_cell_margins(cell, top=130, start=110, bottom=130, end=110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 7) else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run("")
        set_run_font(run, size=9, color="20242B")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("记录后建议双方各自先写事实，再讨论需求和下一步。")
set_run_font(footer_run, size=8, color="7A8494")

doc.save(OUT)
print(OUT.resolve())
